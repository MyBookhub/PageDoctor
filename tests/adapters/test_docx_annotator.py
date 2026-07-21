import io
import zipfile

import docx
import pytest

from pagedoctor.adapters.docx.annotator import COMMENT_AUTHOR, annotate_docx
from pagedoctor.domain.errors import OutputCopyError
from pagedoctor.domain.models.finding import Category, Finding, Priority, Suggestion


def _finding(
    original: str = "Der Hund schläft.",
    proposed: str = "Der Hund schläft tief.",
    reason: str = "Präzisere Formulierung.",
) -> Finding:
    return Finding(
        suggestion=Suggestion(original_text=original, proposed_change=proposed, reason_de=reason),
        category=Category.PROOFREADING,
        priority=Priority.FEHLER,
    )


def _docx(*paragraphs: str) -> bytes:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _comments_xml(content: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return archive.read("word/comments.xml").decode("utf-8")


def _document_xml(content: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_comment_is_anchored_with_author_and_body() -> None:
    source = _docx("Einleitung.", "Der Hund schläft. Danach wacht er auf.")

    annotated = annotate_docx(source, [_finding()])

    comments = _comments_xml(annotated)
    assert COMMENT_AUTHOR in comments
    assert "Der Hund schläft." in comments
    assert "Der Hund schläft tief." in comments
    assert "Präzisere Formulierung." in comments
    document = _document_xml(annotated)
    assert "commentRangeStart" in document
    assert "commentRangeEnd" in document


def test_one_comment_per_finding() -> None:
    source = _docx("Der Hund schläft.", "Die Katze schnurrt laut.")

    annotated = annotate_docx(
        source,
        [_finding(), _finding(original="Die Katze schnurrt", proposed="Die Katze schnurrt leise")],
    )

    assert _comments_xml(annotated).count("w:comment ") == 2


def test_unlocatable_quote_still_lands_as_comment() -> None:
    source = _docx("Ganz anderer Text ohne das Zitat.")

    annotated = annotate_docx(source, [_finding(original="Nicht im Dokument.")])

    comments = _comments_xml(annotated)
    assert "Nicht im Dokument." in comments
    assert "commentRangeStart" in _document_xml(annotated)


def test_document_without_any_run_raises() -> None:
    source = _docx()

    with pytest.raises(OutputCopyError):
        annotate_docx(source, [_finding()])


def test_no_findings_returns_valid_docx_without_comments() -> None:
    source = _docx("Sauberes Manuskript.")

    annotated = annotate_docx(source, [])

    with zipfile.ZipFile(io.BytesIO(annotated)) as archive:
        assert "word/document.xml" in archive.namelist()
        assert "word/comments.xml" not in archive.namelist()


def test_source_text_is_never_modified() -> None:
    source = _docx("Der Hund schläft. Danach wacht er auf.")

    annotated = annotate_docx(source, [_finding()])

    document = docx.Document(io.BytesIO(annotated))
    assert document.paragraphs[0].text == "Der Hund schläft. Danach wacht er auf."
